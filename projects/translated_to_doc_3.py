from deep_translator import GoogleTranslator
from docx import Document


def translate_doc(doc_path: str, target: str = "fr", sections=10):
    doc = Document(doc_path)
    translator = GoogleTranslator(source='auto', target=target)
    print("Translation started...")

    # Collect all text elements for batch translation
    text_elements = []
    element_map = []  # Stores references to original elements for applying translated text

    # Collect paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text = ''.join(run.text for run in paragraph.runs)
            text_elements.append(full_text)
            element_map.append(('paragraph', paragraph))

    # Collect table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        full_text = ''.join(run.text for run in paragraph.runs)
                        text_elements.append(full_text)
                        element_map.append(('cell', paragraph))

    # Translate all collected text at once to improve performance
    translated_texts = translator.translate_batch(text_elements)

    # Apply translations to respective elements
    total_elements = len(element_map)
    elements_per_section = max(1, total_elements // sections)  # Divide into sections for progress

    for index, ((element_type, element), translated_text) in enumerate(zip(element_map, translated_texts)):
        if element_type == 'paragraph' or element_type == 'cell':
            apply_translated_text_to_runs(element, translated_text)

        # Print progress updates based on the section
        if (index + 1) % elements_per_section == 0:
            print(f"Approximately {(index + 1) // elements_per_section * 10}% of document translated.")

    print("Translation completed.")
    doc.save("translated_document_new_code.docx")


def apply_translated_text_to_runs(paragraph, translated_text):
    """Applies translated text to a paragraph while preserving original run styles and ensuring full translation."""
    # Clear all runs except the first to handle the full translation at once
    if paragraph.runs:
        paragraph.runs[0].text = translated_text  # Apply the full translation to the first run
        paragraph.runs[0].font.size = paragraph.runs[0].font.size
        paragraph.runs[0].bold = paragraph.runs[0].bold
        paragraph.runs[0].italic = paragraph.runs[0].italic
        paragraph.runs[0].underline = paragraph.runs[0].underline
        paragraph.runs[0].font.color.rgb = paragraph.runs[0].font.color.rgb
        paragraph.runs[0].font.name = paragraph.runs[0].font.name

        # Clear remaining runs
        for run in paragraph.runs[1:]:
            run.text = ""

