from PyPDF2 import PdfReader, PdfWriter
from docx import Document


def remove_blank_pages(input_pdf, output_pdf):
    reader = PdfReader(input_pdf)
    writer = PdfWriter()

    for page in reader.pages:
        text = page.extract_text()

        if text.strip():
            writer.add_page(page)

    with open(output_pdf, 'wb') as f:
        writer.write(f)

    print(f"Blank pages removed. New PDF saved as: {output_pdf}")


def remove_blank_docx(input_docx, output_docx):
    doc = Document(input_docx)
    new_doc = Document()

    for para in doc.paragraphs:
        if para.text.strip():  # If the paragraph is not blank
            new_paragraph = new_doc.add_paragraph()
            new_paragraph.alignment = para.alignment
            new_paragraph.style = para.style

            # Copy runs to retain inline styling
            for run in para.runs:
                new_run = new_paragraph.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.size = run.font.size
                new_run.font.name = run.font.name
                if run.font.color:
                    new_run.font.color.rgb = run.font.color.rgb

    # Copy tables and other non-paragraph elements like images
    for element in doc.element.body:
        if element.tag.endswith("tbl") or element.tag.endswith("drawing"):
            new_doc.element.body.append(element)

    # Save the new document without blank pages
    new_doc.save(output_docx)
    print(f"Blank pages removed. New document saved as: {output_docx}")
