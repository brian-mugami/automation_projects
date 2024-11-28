import zipfile

from docx import Document
from lxml import etree


def read_word_document(file_path):
    data = {
        "headings": [],
        "body": [],
        "tables": [],
        "headers": [],
        "footers": [],
        "text_boxes": []
    }

    # --- Extract body text and tables ---
    doc = Document(file_path)

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):  # Check if the style is a heading
            data["headings"].append((para.style.name, para.text))  # Save heading level and text
        elif para.text:  # For normal body text
            data["body"].append(para.text)
    # Tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        data["tables"].append(table_data)

    # Headers and footers
    for section in doc.sections:
        if section.header:
            data["headers"].append("\n".join([para.text for para in section.header.paragraphs if para.text]))
        if section.footer:
            data["footers"].append("\n".join([para.text for para in section.footer.paragraphs if para.text]))

    # --- Extract text boxes ---
    with zipfile.ZipFile(file_path) as docx:
        xml_content = docx.read("word/document.xml")
    tree = etree.XML(xml_content)

    # Namespaces for Word documents
    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    # Find all text in text boxes (w:txbxContent)
    text_boxes = tree.findall('.//w:txbxContent', namespaces)
    for box in text_boxes:
        paragraphs = box.findall('.//w:p', namespaces)
        for para in paragraphs:
            texts = para.findall('.//w:t', namespaces)
            data["text_boxes"].append("".join([t.text for t in texts if t.text]))

    return data


file_path = "ERP_Document_Final.docx"  # Replace with your file path
document_data = read_word_document(file_path)

print("Headings:")
for level, heading in document_data["headings"]:
    print(f"{level}: {heading}")

print("\nBody Text:")
print("\n".join(document_data["body"]))

print("\nTables:")
for table in document_data["tables"]:
    for row in table:
        print(row)

print("\nHeaders:")
print("\n".join(document_data["headers"]))

print("\nFooters:")
print("\n".join(document_data["footers"]))

print("\nText Boxes:")
print("\n".join(document_data["text_boxes"]))
