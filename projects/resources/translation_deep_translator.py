import time

import nltk
from deep_translator import GoogleTranslator
from docx import Document
from nltk.tokenize import sent_tokenize

nltk.download('punkt')


def translate_word_document_with_deep_translator(doc_path: str, translated_path: str, target: str = "fr",
                                                 retries: int = 5, delay: int = 2):
    try:
        doc = Document(doc_path)
    except Exception as e:
        yield f"Error loading document: {e}"
        return

    translator = GoogleTranslator(source='auto', target=target)
    yield "Translation started with deep translator"

    elements_translated = 0
    total_elements = sum(1 for _ in iterate_text_elements(doc))
    if total_elements == 0:
        yield "No translatable text found."
        return

    yield f"Found {total_elements} elements to translate."

    for element_type, element in iterate_text_elements(doc):
        original_text = ''.join(run.text for run in element.runs if run.text).strip()
        if not original_text:
            continue

        translated_text = retry_translation(translator, original_text, retries, delay)
        if translated_text:
            apply_translated_text_with_styling(element, translated_text)
        else:
            yield f"Warning: Translation failed for element: {original_text[:50]}..."
            apply_translated_text_with_styling(element, "[Translation Failed]")

        elements_translated += 1
        yield f"Translation progress: {elements_translated}/{total_elements} completed ({(elements_translated / total_elements) * 100:.2f}%)"

    yield "Translation completed."
    try:
        doc.save(translated_path)
        yield f"Translated document saved at: {translated_path}"
    except Exception as e:
        yield f"Error saving translated document: {e}"


def iterate_text_elements(doc):
    """Generator to iterate over all paragraphs and table cells in a Word document."""
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            yield ('paragraph', paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        yield ('cell', paragraph)


def retry_translation(translator, text, retries=5, delay=2):
    """Attempts to translate text with retries."""
    translated_chunks = []
    sentences = sent_tokenize(text)

    for sentence in sentences:
        translated_sentence = None
        for attempt in range(retries):
            try:
                translated_sentence = translator.translate(sentence)
                if translated_sentence is not None:
                    break
            except Exception as e:
                print(f"Error translating sentence on attempt {attempt + 1}: {e}")
                if attempt < retries - 1:
                    time.sleep(delay)

        if translated_sentence:
            translated_chunks.append(translated_sentence)
        else:
            print(f"Failed to translate sentence: {sentence[:50]}...")
            translated_chunks.append("[Translation Failed]")

    return ' '.join(translated_chunks)


def apply_translated_text_with_styling(paragraph, translated_text):
    sentences = sent_tokenize(translated_text)
    original_runs = paragraph.runs

    for run in original_runs:
        run.text = ""

    for i, sentence in enumerate(sentences):
        if i < len(original_runs):
            run = original_runs[i]
        else:
            run = paragraph.add_run()

        run.text = sentence
        if not run.font.name:
            run.font.name = "Times New Roman"
        run.font.size = run.font.size or paragraph.style.font.size
        run.bold = run.bold or False
        run.italic = run.italic or False


def clean_text(text):
    """Cleans up text for consistent formatting."""
    return text.replace('\xa0', ' ').strip()

# import time
#
# import nltk
# from deep_translator import GoogleTranslator
# from docx import Document
# from nltk.tokenize import sent_tokenize
#
# nltk.download('punkt')
# nltk.download('punkt_tab')
#
#
# def translate_word_document_with_deep_translator(doc_path: str, translated_path: str, target: str = "fr",
#                                                  retries: int = 5, delay: int = 2):
#     try:
#         doc = Document(doc_path)
#     except Exception as e:
#         yield f"Error loading document: {e}"
#         return
#
#     translator = GoogleTranslator(source='auto', target=target)
#     yield "Translation started with deep translator"
#
#     elements_translated = 0
#     total_elements = sum(1 for _ in iterate_text_elements(doc))
#     if total_elements == 0:
#         yield "No translatable text found."
#         return
#
#     yield f"Found {total_elements} elements to translate."
#
#     for element_type, element in iterate_text_elements(doc):
#         original_text = clean_text(''.join(run.text for run in element.runs if run.text).strip())
#         if not original_text:
#             continue
#         translated_text = retry_translation(translator, original_text, retries, delay)
#         if translated_text:
#             apply_translated_text(element, translated_text)
#         else:
#             yield f"Warning: Translation failed for element: {original_text[:50]}..."
#             apply_translated_text(element, "[Translation Failed]")
#
#         elements_translated += 1
#         yield f"Translation progress: {elements_translated}/{total_elements} completed ({(elements_translated / total_elements) * 100:.2f}%)"
#
#     yield "Translation completed."
#     try:
#         doc.save(translated_path)
#         yield f"Translated document saved at: {translated_path}"
#     except Exception as e:
#         yield f"Error saving translated document: {e}"
#
#
# def iterate_text_elements(doc):
#     """Generator to iterate over all paragraphs and table cells in a Word document."""
#     for paragraph in doc.paragraphs:
#         if paragraph.text.strip():
#             yield ('paragraph', paragraph)
#
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for paragraph in cell.paragraphs:
#                     if paragraph.text.strip():
#                         yield ('cell', paragraph)
#
#
# def retry_translation(translator, text, retries=5, delay=2):
#     translated_chunks = []
#     sentences = sent_tokenize(text)
#
#     for sentence in sentences:
#         translated_sentence = None
#         for attempt in range(retries):
#             try:
#                 translated_sentence = translator.translate(sentence)
#                 if translated_sentence is not None:
#                     break
#             except Exception as e:
#                 print(f"Error translating sentence on attempt {attempt + 1}: {e}")
#                 if attempt < retries - 1:
#                     time.sleep(delay)
#
#         if translated_sentence:
#             translated_chunks.append(translated_sentence)
#         else:
#             print(f"Failed to translate sentence: {sentence[:50]}...")
#             translated_chunks.append("[Translation Failed]")  # Placeholder for untranslatable text
#
#     return ' '.join(translated_chunks)
#
#
# def apply_translated_text(paragraph, translated_text):
#     for run in paragraph.runs:
#         run.text = ""
#     paragraph.add_run(translated_text)
#
#
# def clean_text(text):
#     return text.replace('\xa0', ' ').strip()
