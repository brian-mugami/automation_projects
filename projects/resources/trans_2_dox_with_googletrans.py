import time
import nltk
from docx import Document
# from googletrans import Translator
from nltk.tokenize import sent_tokenize

nltk.download('punkt')


def translate_word_document(doc_path: str, translated_path: str, target: str = "fr", retries: int = 5, delay: int = 2):
    try:
        doc = Document(doc_path)
    except Exception as e:
        yield f"Error loading document: {e}"
        return
    # translator = Translator()
    print("started")
    yield "data: Translation started...\n\n"

    elements_translated = 0
    total_elements = sum(1 for _ in iterate_text_elements(doc))
    if total_elements == 0:
        yield "data: No translatable text found.\n\n"
        return

    yield f"data: Found {total_elements} elements to translate.\n\n"

    for element_type, element in iterate_text_elements(doc):
        for run in element.runs:
            original_text = clean_text(run.text.strip())
            if not original_text:
                continue  # Skip empty text

            translated_text = retry_translation(translator, original_text, target, retries, delay)
            if translated_text:
                apply_translated_text(run, translated_text)
            else:
                yield f"data: Warning: Translation failed for text: {original_text[:50]}...\n\n"
                apply_translated_text(run, "[Translation Failed]")

            elements_translated += 1

        progress = f"Translation progress: {elements_translated}/{total_elements} completed ({(elements_translated / total_elements) * 100:.2f}%)"
        yield f"data: {progress}\n\n"

    yield "data: Translation completed.\n\n"
    try:
        doc.save(translated_path)
        yield f"data: Translated document saved at: {translated_path}\n\n"
    except Exception as e:
        yield f"data: Error saving translated document: {e}\n\n"


def iterate_text_elements(doc):
    """Generator to iterate over all paragraphs and table cells in a Word document."""
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            yield paragraph

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        yield paragraph


def retry_translation(translator, text, target, retries=5, delay=2):
    sentences = sent_tokenize(text)
    translated_chunks = []

    for sentence in sentences:
        for attempt in range(retries):
            try:
                translated_sentence = translator.translate(sentence, dest=target).text
                translated_chunks.append(translated_sentence)
                break
            except Exception as e:
                print(f"Error translating sentence on attempt {attempt + 1}: {e}")
                if attempt < retries - 1:
                    time.sleep(delay)
                else:
                    translated_chunks.append("[Translation Failed]")
    return ' '.join(translated_chunks)


def apply_translated_text(run, translated_text):
    """Preserve original styling and apply translated text."""
    font_name = run.font.name or "Times New Roman"
    font_size = run.font.size
    is_bold = run.bold
    is_italic = run.italic
    is_underline = run.underline

    # Clear the existing text
    run.text = ""
    new_run = run._element  # Replace run content
    new_run.text = translated_text

    # Apply styling
    run.font.name = font_name
    if font_size:
        run.font.size = font_size
    run.bold = is_bold
    run.italic = is_italic
    run.underline = is_underline


def clean_text(text):
    return text.replace('\xa0', ' ').strip()


# import time
#
# import nltk
# from docx import Document
# from googletrans import Translator
# from nltk.tokenize import sent_tokenize
#
# nltk.download('punkt')
#
#
# def translate_word_document(doc_path: str, translated_path: str, target: str = "fr", retries: int = 5, delay: int = 2):
#     try:
#         doc = Document(doc_path)
#     except Exception as e:
#         yield f"Error loading document: {e}"
#         return
#     translator = Translator()
#     print("started")
#     yield "data: Translation started...\n\n"
#
#     elements_translated = 0
#     total_elements = sum(1 for _ in iterate_text_elements(doc))
#     if total_elements == 0:
#         yield "data: No translatable text found.\n\n"
#         return
#
#     yield f"data: Found {total_elements} elements to translate.\n\n"
#
#     for element_type, element in iterate_text_elements(doc):
#         original_text = clean_text(''.join(run.text for run in element.runs if run.text).strip())
#         if not original_text:
#             continue  # Skip empty text
#
#         translated_text = retry_translation(translator, original_text, target, retries, delay)
#         if translated_text:
#             apply_translated_text(element, translated_text)
#         else:
#             yield f"data: Warning: Translation failed for element: {original_text[:50]}...\n\n"
#             apply_translated_text(element, "[Translation Failed]")
#
#         elements_translated += 1
#         progress = f"Translation progress: {elements_translated}/{total_elements} completed ({(elements_translated / total_elements) * 100:.2f}%)"
#         yield f"data: {progress}\n\n"
#
#     yield "data: Translation completed.\n\n"
#     try:
#         doc.save(translated_path)
#         yield f"data: Translated document saved at: {translated_path}\n\n"
#     except Exception as e:
#         yield f"data: Error saving translated document: {e}\n\n"
#
#
# def iterate_text_elements(doc):
#     """Generator to iterate over all paragraphs and table cells in a Word document."""
#     for paragraph in doc.paragraphs:
#         if paragraph.text.strip():
#             yield ('paragraph', paragraph)
#
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for paragraph in cell.paragraphs:
#                     if paragraph.text.strip():
#                         yield ('cell', paragraph)
#
#
# def retry_translation(translator, text, target, retries=5, delay=2):
#     translated_chunks = []
#     sentences = sent_tokenize(text)  # Tokenize by sentence
#
#     for sentence in sentences:
#         for attempt in range(retries):
#             try:
#                 translated_chunks.append(translator.translate(sentence, dest=target).text)
#                 break  # Exit retry loop for this sentence on success
#             except Exception as e:
#                 print(f"Error translating sentence on attempt {attempt + 1}: {e}")
#                 if attempt < retries - 1:
#                     time.sleep(delay)
#                 else:
#                     print(f"Failed to translate sentence: {sentence[:50]}...")
#                     translated_chunks.append("[Translation Failed]")  # Placeholder for untranslatable text
#
#     return ' '.join(translated_chunks)
#
#
# def apply_translated_text(paragraph, translated_text):
#     for run in paragraph.runs:
#         run.text = ""  # Clear existing text
#     paragraph.add_run(translated_text)  # Add the translated text as a single run
#
#
# def clean_text(text):
#     return text.replace('\xa0', ' ').strip()
