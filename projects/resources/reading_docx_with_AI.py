import zipfile

from docx import Document
from lxml import etree


def read_word_document(file_path):
    data = {
        "headings": [],
        "body": [],
        "tables": [],
        "headers": [],
        "footers": [],
        "text_boxes": []
    }
    doc = Document(file_path)

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            data["headings"].append((para.style.name, para.text))
        elif para.text:
            data["body"].append(para.text)
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        data["tables"].append(table_data)

    for section in doc.sections:
        if section.header:
            data["headers"].append("\n".join([para.text for para in section.header.paragraphs if para.text]))
        if section.footer:
            data["footers"].append("\n".join([para.text for para in section.footer.paragraphs if para.text]))

    with zipfile.ZipFile(file_path) as docx:
        xml_content = docx.read("word/document.xml")
    tree = etree.XML(xml_content)

    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    text_boxes = tree.findall('.//w:txbxContent', namespaces)
    for box in text_boxes:
        paragraphs = box.findall('.//w:p', namespaces)
        for para in paragraphs:
            texts = para.findall('.//w:t', namespaces)
            data["text_boxes"].append("".join([t.text for t in texts if t.text]))

    return data


def split_data_content(data, specified_area, max_length=6000):
    area = data[specified_area]
    return [
        area[i:i + max_length] for i in range(0, len(area), max_length)
    ]
