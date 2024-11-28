import time

from deep_translator import GoogleTranslator
from docx import Document


def translate_text_individually(doc_path: str, target: str = "fr", retries: int = 5):
    # Open the document
    doc = Document(doc_path)
    translator = GoogleTranslator(source='auto', target=target)
    print("Translation started...")

    elements_translated = 0
    total_elements = sum(1 for _ in iterate_text_elements(doc))
    if total_elements == 0:
        print("No translatable text found.")
        return

    print(f"Found {total_elements} elements to translate.")

    for element_type, element in iterate_text_elements(doc):
        original_text = ''.join(run.text for run in element.runs if run.text).strip()
        if not original_text:
            continue  # Skip empty text

        # Retry mechanism for each element
        translated_text = retry_individual_translation(translator, original_text, retries)
        if translated_text:
            apply_translated_text_to_runs(element, translated_text)
        else:
            print("Warning: Translation failed for an element, keeping original text.")

        elements_translated += 1
        print_progress(elements_translated, total_elements)

    print("Translation completed.")
    doc.save("translated_document_100.docx")


def iterate_text_elements(doc):
    """Generator to iterate over all paragraphs and table cells in a Word document."""
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            yield ('paragraph', paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        yield ('cell', paragraph)


def retry_individual_translation(translator, text, retries=5, delay=2):
    """Attempts to translate a single text with retries."""
    for attempt in range(retries):
        try:
            return translator.translate(text)
        except Exception as e:
            print(f"Error translating text on attempt {attempt + 1}: {e}")
            if attempt < retries - 1:
                time.sleep(delay)  # Wait before retrying
            else:
                print("Translation failed after multiple attempts.")
                return None


def apply_translated_text_to_runs(paragraph, translated_text):
    """Applies translated text to a paragraph while preserving original run styles."""
    translated_words = translated_text.split()
    current_word_index = 0

    for run in paragraph.runs:
        if run.text:
            run_text_words = run.text.split()
            words_to_take = translated_words[current_word_index:current_word_index + len(run_text_words)]
            run.text = ' '.join(words_to_take)
            current_word_index += len(words_to_take)


def print_progress(completed, total):
    """Prints translation progress incrementally."""
    print(f"Translation progress: {completed}/{total} completed ({(completed / total) * 100:.2f}%)")


