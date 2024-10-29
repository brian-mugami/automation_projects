from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from openpyxl import load_workbook


def optimize_column_widths(sheet):
    """Optimize column widths based on the max length of the data, ignoring merged cells."""
    column_widths = {}
    for column in sheet.columns:
        max_length = 0
        column_letter = None

        # Find the first non-merged cell in the column to get the column letter
        for cell in column:
            if not any(cell.coordinate in merged_range for merged_range in sheet.merged_cells.ranges):
                column_letter = cell.column_letter  # Get the column letter
                break

        # Calculate max length of non-merged cells in the column
        for cell in column:
            if not any(cell.coordinate in merged_range for merged_range in sheet.merged_cells.ranges):
                if cell.value is not None:
                    max_length = max(max_length, len(str(cell.value)))

        if column_letter:
            adjusted_width = max_length + 2
            column_widths[column_letter] = Inches(adjusted_width / 10)

    return column_widths


def is_sheet_content_present(sheet):
    """Check if the sheet contains any non-empty cells."""
    for row in sheet.iter_rows(values_only=True):
        if any(cell is not None for cell in row):
            return True
    return False


def apply_cell_formatting(cell, excel_cell):
    """Apply formatting from Excel cell to Word cell."""
    cell.text = str(excel_cell.value) if excel_cell.value is not None else ""
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell.paragraphs[0].runs[0].font.size = Pt(10)  # Default font size

    # Apply background color if available
    if excel_cell.fill.start_color.index != '00000000':  # Check for a non-default color
        try:
            rgb_color = RGBColor(int(excel_cell.fill.start_color.index, 16),
                                 int(excel_cell.fill.start_color.index, 16),
                                 int(excel_cell.fill.start_color.index, 16))
            cell._element.get_or_add('w:shd').set('w:fill', rgb_color)
        except ValueError:
            print(f"Invalid background color index: {excel_cell.fill.start_color.index}")

    # Apply font color
    if excel_cell.font.color and excel_cell.font.color.index is not None:
        if isinstance(excel_cell.font.color.index, str):  # Ensure it's a string before conversion
            try:
                color_value = int(excel_cell.font.color.index, 16)
                r = (color_value >> 16) & 0xFF
                g = (color_value >> 8) & 0xFF
                b = color_value & 0xFF
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(r, g, b)
            except ValueError:
                print(f"Invalid font color index: {excel_cell.font.color.index}")

    # Set borders
    set_cell_borders(cell, excel_cell)


def set_cell_borders(cell, excel_cell):
    """Set cell borders based on the Excel cell's borders."""
    # Example: Add code to set borders if needed
    # Borders in docx can be set using styles, but is more complex than Excel.
    # This is a placeholder for border implementation.
    pass


def convert_excel_to_word(excel_path, output_word, orientation='Landscape'):
    workbook = load_workbook(excel_path)
    document = Document()

    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]

        if not is_sheet_content_present(sheet):
            print(f"Skipping empty sheet: {sheet_name}")
            continue

        # Set section orientation
        section = document.sections[-1]
        if orientation == 'Landscape':
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width

        # Add sheet title
        document.add_heading(sheet_name, level=1)

        # Optimize column widths
        column_widths = optimize_column_widths(sheet)

        # Create table
        table = document.add_table(rows=1, cols=sheet.max_column)
        table.autofit = False

        header_row = table.rows[0]
        for idx, cell in enumerate(sheet[1]):
            if not any(cell.coordinate in merged_range for merged_range in
                       sheet.merged_cells.ranges):  # Skip merged cells in headers
                header_cell = header_row.cells[idx]
                header_cell.text = str(cell.value) if cell.value else ""
                header_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                header_cell.paragraphs[0].runs[0].bold = True
                header_cell.paragraphs[0].runs[0].font.size = Pt(10)
                table.columns[idx].width = column_widths[cell.column_letter]

        # Add data rows
        for row in sheet.iter_rows(min_row=2, values_only=False):
            word_row = table.add_row().cells
            for idx, excel_cell in enumerate(row):
                if not any(excel_cell.coordinate in merged_range for merged_range in
                           sheet.merged_cells.ranges):  # Skip merged cells
                    apply_cell_formatting(word_row[idx], excel_cell)

        # Add page break between sheets
        if sheet_name != workbook.sheetnames[-1]:
            document.add_page_break()

    # Save the Word document
    document.save(output_word)
    print(f"Converted Excel workbook to Word: {output_word}")
