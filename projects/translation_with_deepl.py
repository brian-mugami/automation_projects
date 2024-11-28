from deepl import Translator
from docx import Document


def translate_doc(doc_path: str, target: str = "fr"):
    doc = Document(doc_path)
    translator = Translator("b1955498-473f-46db-8189-41537a41577a:fx")  # Replace with your Deepl API key
    print("Translation started...")

    total_elements = len(doc.paragraphs) + sum(len(row.cells) for table in doc.tables for row in table.rows)
    completed_elements = 0

    # Translate paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            try:
                full_text = ''.join(run.text for run in paragraph.runs)
                translated_text = translator.translate_text(full_text, target_lang=target).text  # Get translated text
                apply_translated_text_to_runs(paragraph, translated_text)
            except Exception as e:
                print(f"Error translating paragraph: {e}")
            completed_elements += 1
            print_progress(completed_elements, total_elements)

    # Translate text in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        try:
                            full_text = ''.join(run.text for run in paragraph.runs)
                            translated_text = translator.translate_text(full_text,
                                                                        target_lang=target).text  # Get translated text
                            apply_translated_text_to_runs(paragraph, translated_text)
                        except Exception as e:
                            print(f"Error translating table cell: {e}")
                    completed_elements += 1
                    print_progress(completed_elements, total_elements)

    print("Translation completed.")
    doc.save("translated_document_15.docx")


def apply_translated_text_to_runs(paragraph, translated_text):
    """Applies translated text to a paragraph while preserving original run styles."""
    translated_words = translated_text.split()
    current_word_index = 0

    for run in paragraph.runs:
        run_text_words = run.text.split()
        words_to_take = translated_words[current_word_index:current_word_index + len(run_text_words)]
        run.text = ' '.join(words_to_take)
        current_word_index += len(words_to_take)

        # Retain original styling for each run
        run.font.size = run.font.size
        run.font.bold = run.font.bold
        run.font.italic = run.font.italic
        run.font.underline = run.font.underline
        run.font.color.rgb = run.font.color.rgb
        run.font.name = run.font.name


def print_progress(completed, total):
    """Prints translation progress in increments of 10%."""
    progress = int((completed / total) * 100)
    if progress % 10 == 0:  # Only print every 10%
        print(f"Translation progress: {progress}% completed")


