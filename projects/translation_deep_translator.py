import time
from deep_translator import GoogleTranslator
from docx import Document
from nltk.tokenize import sent_tokenize
import nltk

# Download punkt for sentence tokenization
nltk.download('punkt')
nltk.download('punkt_tab')


def translate_word_document(doc_path: str, translated_path: str, target: str = "fr", retries: int = 5, delay: int = 2):
    try:
        doc = Document(doc_path)
    except Exception as e:
        print(f"Error loading document: {e}")
        return

    translator = GoogleTranslator(source='auto', target=target)
    print("Translation started...")

    elements_translated = 0
    total_elements = sum(1 for _ in iterate_text_elements(doc))
    if total_elements == 0:
        print("No translatable text found.")
        return

    print(f"Found {total_elements} elements to translate.")

    for element_type, element in iterate_text_elements(doc):
        original_text = clean_text(''.join(run.text for run in element.runs if run.text).strip())
        if not original_text:
            continue  # Skip empty text

        # Retry mechanism for each element
        translated_text = retry_translation(translator, original_text, retries, delay)
        if translated_text:
            apply_translated_text(element, translated_text)
        else:
            print(f"Warning: Translation failed for element: {original_text[:50]}...")
            apply_translated_text(element, "[Translation Failed]")

        elements_translated += 1
        print_progress(elements_translated, total_elements)

    print("Translation completed.")
    try:
        doc.save(translated_path)
        print(f"Translated document saved at: {translated_path}")
    except Exception as e:
        print(f"Error saving translated document: {e}")


def iterate_text_elements(doc):
    """Generator to iterate over all paragraphs and table cells in a Word document."""
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            yield ('paragraph', paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        yield ('cell', paragraph)


def retry_translation(translator, text, retries=5, delay=2):
    """
    Attempts to translate a text with retries. Splits text into sentences for better translation quality.
    """
    translated_chunks = []
    sentences = sent_tokenize(text)  # Tokenize by sentence

    for sentence in sentences:
        translated_sentence = None
        for attempt in range(retries):
            try:
                # Translate the sentence using `deep_translator`
                translated_sentence = translator.translate(sentence)
                if translated_sentence is not None:
                    break  # Exit retry loop for this sentence on success
            except Exception as e:
                print(f"Error translating sentence on attempt {attempt + 1}: {e}")
                if attempt < retries - 1:
                    time.sleep(delay)

        if translated_sentence:
            translated_chunks.append(translated_sentence)
        else:
            print(f"Failed to translate sentence: {sentence[:50]}...")
            translated_chunks.append("[Translation Failed]")  # Placeholder for untranslatable text

    return ' '.join(translated_chunks)



def apply_translated_text(paragraph, translated_text):
    """
    Applies translated text to a paragraph, clearing existing runs and replacing with the translated text.
    """
    for run in paragraph.runs:
        run.text = ""  # Clear existing text
    paragraph.add_run(translated_text)  # Add the translated text as a single run


def clean_text(text):
    """Cleans text by removing unwanted characters."""
    return text.replace('\xa0', ' ').strip()


def print_progress(completed, total):
    """Prints the progress of translation."""
    print(f"Translation progress: {completed}/{total} completed ({(completed / total) * 100:.2f}%)")


# Example usage
input_file = "EOI notice in English.docx"
output_file = "EOI notice in English_trans.docx"
translate_word_document(input_file, output_file, target="fr")
